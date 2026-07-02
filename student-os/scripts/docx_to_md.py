#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path

from xml.etree.ElementTree import QName


def load_document():
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit(
            "Missing dependency 'python-docx'. Install the packages from requirements.txt before running this script."
        ) from exc
    return Document


def markdown_cell(text: str) -> str:
    return text.replace("|", "\\|").replace("\n", "<br>")


def iter_block_items(document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == QName(
            "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "p",
        ).text:
            yield "paragraph", child
        elif child.tag == QName(
            "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "tbl",
        ).text:
            yield "table", child


def paragraph_text(document, element) -> str:
    for para in document.paragraphs:
        if para._element is element:
            return para.text.strip()
    return ""


def table_rows(document, element) -> list[list[str]]:
    for table in document.tables:
        if table._element is element:
            return [
                [markdown_cell(cell.text.strip()) for cell in row.cells]
                for row in table.rows
            ]
    return []


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert a DOCX file into a markdown reference draft.")
    parser.add_argument("docx", help="Path to the DOCX file")
    parser.add_argument("--output", required=True, help="Output markdown path")
    args = parser.parse_args()

    docx_path = Path(args.docx).resolve()
    output_path = Path(args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    Document = load_document()
    doc = Document(str(docx_path))
    lines = [
        "---",
        "type: imported-reference",
        "course:",
        "status: active",
        "created:",
        "updated:",
        "tags: [import, reference]",
        f"source_file: {docx_path}",
        "import_method: docx-to-md",
        "repair_status:",
        "derived_from_import:",
        "---",
        "",
        f"# Imported Reference - {docx_path.stem}",
        "",
        "## Source",
        "",
        f"- Source file: {docx_path}",
        "- Import method: docx-to-md",
        "",
        "## Imported Content",
        "",
    ]

    table_index = 0
    for block_type, element in iter_block_items(doc):
        if block_type == "paragraph":
            text = paragraph_text(doc, element)
            if text:
                lines.extend([text, ""])
            continue

        table_index += 1
        rows = table_rows(doc, element)
        lines.extend([f"## Table {table_index}", ""])
        if rows:
            header = rows[0]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in rows[1:]:
                values = row[: len(header)] + ([""] * max(0, len(header) - len(row)))
                lines.append("| " + " | ".join(values) + " |")
            lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps({"output": str(output_path)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
